


print("hello")
import docx

doc = docx.Document('example.docx')

# количество абзацев в документе
print(len(doc.paragraphs))

# текст первого абзаца в документе
print(doc.paragraphs[0].text)

doc.add_paragraph('Здравствуй, мир!')

# добавляем еще два параграфа
par1 = doc.add_paragraph('Это второй абзац.')
par2 = doc.add_paragraph('Это третий абзац.')

# добавляем текст во второй параграф
par1.add_run(' Этот текст был добавлен во второй абзац.')

# добавляем текст в третий параграф
par2.add_run(' Добавляем текст в третий абзац.').bold = True



doc.save('example.docx')
